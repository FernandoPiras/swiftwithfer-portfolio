#!/usr/bin/env python3
"""
SwiftWithFer Discovery Document Generator
Produces: PDF, Fillable PDF (AcroForm), DOCX, Markdown, JSON
"""

import json
import os
import sys
from datetime import datetime
from pathlib import Path

# Add scripts dir to path
SCRIPT_DIR = Path(__file__).parent
ROOT = SCRIPT_DIR.parent
OUTPUT = ROOT / "output"
ASSETS = ROOT / "assets"
LOGO = ASSETS / "logo.png"
PUBLIC_DISCOVERY = ROOT.parent / "public" / "documents" / "discovery"

sys.path.insert(0, str(SCRIPT_DIR))
from discovery_schema import META, SECTIONS, SECTORS, get_schema, iter_fields

# ─── Design tokens ───────────────────────────────────────────────────────────
COLORS = {
    "primary": "#0066FF",
    "primary_dark": "#0044CC",
    "accent": "#00D4FF",
    "dark": "#0A0A0F",
    "text": "#1A1A2E",
    "text_light": "#6B7280",
    "border": "#E5E7EB",
    "bg_light": "#F8FAFC",
    "bg_section": "#F1F5F9",
    "white": "#FFFFFF",
}


def hex_to_rgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i : i + 2], 16) / 255 for i in (0, 2, 4))


# ─── Markdown ────────────────────────────────────────────────────────────────
def field_md(field):
    fid = field["id"]
    label = field["label"]
    ftype = field["type"]
    req = " *(obbligatorio)*" if field.get("required") else ""

    if ftype == "checkbox":
        return f"- [ ] **{label}**{req}\n"
    if ftype == "checkbox_group":
        lines = [f"**{label}**{req}\n"]
        for opt in field.get("options", []):
            lines.append(f"- [ ] {opt}\n")
        return "\n".join(lines) + "\n"
    if ftype == "select":
        opts = " · ".join(field.get("options", []))
        return f"**{label}**{req}\n\n> Opzioni: {opts}\n\n`{fid}`: \n\n"
    if ftype == "textarea":
        return f"**{label}**{req}\n\n```\n{fid}:\n\n\n\n\n```\n\n"
    return f"**{label}**{req}\n\n`{fid}`: \n\n"


def generate_markdown():
    lines = [
        f"# {META['title']}\n",
        f"*{META['subtitle']}*\n",
        "\n> **Come compilare:** compila ogni campo sotto. I campi con * sono obbligatori. "
        "Per lo sviluppo con AI, usa il file JSON e genera il prompt con "
        "`python3 discovery-document/scripts/generate_cursor_prompt.py`.\n",
        "\n---\n",
        f"| Versione | {META['version']} |",
        f"| Autore | {META['author']} |",
        f"| Data | _____________ |",
        f"| Cliente | _____________ |",
        f"| Progetto | _____________ |",
        f"| Confidenzialità | {META['confidentiality']} |\n",
        "\n---\n",
        "## Indice\n",
    ]
    for s in SECTIONS:
        lines.append(f"{s['number']}. [{s['title']}](#{s['id']})\n")

    lines.append("\n---\n")

    for s in SECTIONS:
        lines.append(f"\n## {s['number']}. {s['title']} {{#{s['id']}}}\n")
        lines.append(f"*{s['description']}*\n")
        for sub in s["subsections"]:
            lines.append(f"\n### {sub['title']}\n")
            for field in sub["fields"]:
                lines.append(field_md(field))

    lines.append("\n---\n")
    lines.append(f"*© {datetime.now().year} SwiftWithFer — {META['website']}*\n")

    path = OUTPUT / "SwiftWithFer-Discovery-Document.md"
    path.write_text("".join(lines), encoding="utf-8")
    print(f"✓ Markdown: {path}")
    return path


# ─── JSON ────────────────────────────────────────────────────────────────────
def generate_json():
    schema = get_schema()
    schema["meta"]["generated_at"] = datetime.now().isoformat()
    schema["meta"]["format"] = "ai-ready"
    schema["meta"]["usage"] = META["instructions"]
    schema["meta"]["cursor_command"] = (
        "python3 discovery-document/scripts/generate_cursor_prompt.py path/to/compilato.json"
    )

    for section in schema["sections"]:
        for subsection in section["subsections"]:
            for field in subsection["fields"]:
                if field.get("generated"):
                    field["value"] = field.get("template", "")
                    field["note"] = (
                        "Auto-generato: sostituire {{field_id}} con i valori compilati"
                    )
                else:
                    field["value"] = None
                    if field["type"] == "checkbox":
                        field["value"] = False
                    elif field["type"] == "checkbox_group":
                        field["value"] = {o: False for o in field.get("options", [])}

    path = OUTPUT / "SwiftWithFer-Discovery-Document.json"
    path.write_text(json.dumps(schema, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"✓ JSON: {path}")
    return path


# ─── DOCX ────────────────────────────────────────────────────────────────────
def generate_docx():
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Helvetica Neue"
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    def add_horizontal_line(paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "0066FF")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # Cover
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(1.2))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(META["title"])
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(META["subtitle"])
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Versione", META["version"]),
        ("Data", "_____________"),
        ("Cliente", "_____________"),
        ("Progetto", "_____________"),
        ("Confidenzialità", META["confidentiality"]),
    ]
    for i, (k, v) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = k
        meta_table.rows[i].cells[1].text = v

    doc.add_page_break()

    # TOC
    toc_title = doc.add_heading("Indice", level=1)
    for s in SECTIONS:
        p = doc.add_paragraph(f"{s['number']}. {s['title']}")
        p.paragraph_format.left_indent = Cm(0.5)

    doc.add_page_break()

    # Sections
    for s in SECTIONS:
        h = doc.add_heading(f"{s['number']}. {s['title']}", level=1)
        desc = doc.add_paragraph(s["description"])
        desc.runs[0].italic = True
        desc.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        add_horizontal_line(desc)

        for sub in s["subsections"]:
            doc.add_heading(sub["title"], level=2)

            if any(f["type"] in ("checkbox", "checkbox_group") for f in sub["fields"]):
                # Mixed: use paragraphs for checkboxes, table for text
                text_fields = [f for f in sub["fields"] if f["type"] not in ("checkbox", "checkbox_group")]
                check_fields = [f for f in sub["fields"] if f["type"] in ("checkbox", "checkbox_group")]

                if text_fields:
                    table = doc.add_table(rows=len(text_fields), cols=2)
                    table.style = "Table Grid"
                    for i, field in enumerate(text_fields):
                        label = field["label"] + (" *" if field.get("required") else "")
                        table.rows[i].cells[0].text = label
                        table.rows[i].cells[0].width = Cm(6)
                        cell = table.rows[i].cells[1]
                        if field["type"] == "textarea":
                            cell.text = "\n\n\n"
                        elif field["type"] == "select":
                            cell.text = " | ".join(field.get("options", []))
                        else:
                            cell.text = ""

                for field in check_fields:
                    if field["type"] == "checkbox":
                        p = doc.add_paragraph(f"☐ {field['label']}")
                    else:
                        p = doc.add_paragraph(field["label"])
                        for opt in field.get("options", []):
                            doc.add_paragraph(f"  ☐ {opt}")
            else:
                table = doc.add_table(rows=len(sub["fields"]), cols=2)
                table.style = "Table Grid"
                for i, field in enumerate(sub["fields"]):
                    label = field["label"] + (" *" if field.get("required") else "")
                    table.rows[i].cells[0].text = label
                    cell = table.rows[i].cells[1]
                    if field["type"] == "textarea":
                        cell.text = "\n\n\n\n"
                    elif field["type"] == "select":
                        cell.text = " | ".join(field.get("options", []))
                    else:
                        cell.text = ""

        doc.add_page_break()

    path = OUTPUT / "SwiftWithFer-Discovery-Document.docx"
    doc.save(str(path))
    print(f"✓ DOCX: {path}")
    return path


# ─── PDF (Professional) ─────────────────────────────────────────────────────
def generate_pdf():
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm, mm
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, Image, KeepTogether,
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    path = OUTPUT / "SwiftWithFer-Discovery-Document.pdf"
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        topMargin=2.5 * cm,
        bottomMargin=2 * cm,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        title=META["title"],
        author=META["author"],
    )

    styles = getSampleStyleSheet()
    primary = colors.HexColor(COLORS["primary"])
    dark = colors.HexColor(COLORS["dark"])
    text_light = colors.HexColor(COLORS["text_light"])
    border = colors.HexColor(COLORS["border"])
    bg_section = colors.HexColor(COLORS["bg_section"])

    s_title = ParagraphStyle("DocTitle", parent=styles["Title"], fontSize=26, textColor=dark, spaceAfter=6, alignment=TA_CENTER)
    s_subtitle = ParagraphStyle("DocSub", parent=styles["Normal"], fontSize=11, textColor=text_light, alignment=TA_CENTER, spaceAfter=20)
    s_h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=18, textColor=primary, spaceBefore=16, spaceAfter=8, borderPadding=4)
    s_h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=13, textColor=dark, spaceBefore=12, spaceAfter=6)
    s_desc = ParagraphStyle("Desc", parent=styles["Normal"], fontSize=9, textColor=text_light, spaceAfter=10)
    s_field_label = ParagraphStyle("FieldLabel", parent=styles["Normal"], fontSize=9, textColor=dark, leading=12)
    s_field_value = ParagraphStyle("FieldValue", parent=styles["Normal"], fontSize=9, textColor=text_light, leading=12)
    s_toc = ParagraphStyle("TOC", parent=styles["Normal"], fontSize=11, textColor=dark, leftIndent=12, spaceAfter=4)
    s_footer = ParagraphStyle("Footer", parent=styles["Normal"], fontSize=7, textColor=text_light)

    story = []
    page_num = [0]

    def on_page(canvas, doc):
        page_num[0] += 1
        canvas.saveState()
        w, h = A4
        # Header line
        canvas.setStrokeColor(primary)
        canvas.setLineWidth(1.5)
        canvas.line(2 * cm, h - 1.5 * cm, w - 2 * cm, h - 1.5 * cm)
        # Logo small
        if LOGO.exists():
            canvas.drawImage(str(LOGO), 2 * cm, h - 1.4 * cm, width=0.8 * cm, height=0.8 * cm, preserveAspectRatio=True, mask="auto")
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(text_light)
        canvas.drawString(3 * cm, h - 1.2 * cm, "SwiftWithFer")
        canvas.drawRightString(w - 2 * cm, h - 1.2 * cm, META["confidentiality"])
        # Footer
        canvas.setStrokeColor(border)
        canvas.setLineWidth(0.5)
        canvas.line(2 * cm, 1.5 * cm, w - 2 * cm, 1.5 * cm)
        canvas.setFont("Helvetica", 7)
        canvas.drawString(2 * cm, 1 * cm, f"© {datetime.now().year} SwiftWithFer — {META['website']}")
        canvas.drawRightString(w - 2 * cm, 1 * cm, f"Pagina {page_num[0]}")
        canvas.restoreState()

    # Cover
    if LOGO.exists():
        img = Image(str(LOGO), width=3 * cm, height=3 * cm)
        img.hAlign = "CENTER"
        story.append(Spacer(1, 3 * cm))
        story.append(img)

    story.append(Spacer(1, 1 * cm))
    story.append(Paragraph(META["title"], s_title))
    story.append(Paragraph(META["subtitle"], s_subtitle))

    cover_data = [
        ["Versione", META["version"]],
        ["Data", "_______________"],
        ["Cliente", "_______________"],
        ["Progetto", "_______________"],
        ["Confidenzialità", META["confidentiality"]],
    ]
    cover_table = Table(cover_data, colWidths=[4 * cm, 10 * cm])
    cover_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("TEXTCOLOR", (0, 0), (0, -1), text_light),
        ("TEXTCOLOR", (1, 0), (1, -1), dark),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("ALIGN", (0, 0), (0, -1), "RIGHT"),
        ("ALIGN", (1, 0), (1, -1), "LEFT"),
    ]))
    story.append(cover_table)
    story.append(PageBreak())

    # TOC
    story.append(Paragraph("Indice", s_h1))
    story.append(Spacer(1, 0.3 * cm))
    for s in SECTIONS:
        story.append(Paragraph(f"{s['number']}. {s['title']}", s_toc))
    story.append(PageBreak())

    # Content
    for s in SECTIONS:
        block = [
            Paragraph(f"{s['number']}. {s['title']}", s_h1),
            Paragraph(s["description"], s_desc),
        ]

        for sub in s["subsections"]:
            block.append(Paragraph(sub["title"], s_h2))

            rows = []
            for field in sub["fields"]:
                label = field["label"]
                if field.get("required"):
                    label += " *"
                ftype = field["type"]

                if ftype == "checkbox":
                    val = "☐"
                elif ftype == "checkbox_group":
                    val = "\n".join(f"☐ {o}" for o in field.get("options", []))
                elif ftype == "select":
                    val = "  ▸  ".join(field.get("options", []))
                elif ftype == "textarea":
                    val = "\n\n\n\n"
                else:
                    val = "_______________________________________________"

                rows.append([
                    Paragraph(label, s_field_label),
                    Paragraph(val.replace("\n", "<br/>"), s_field_value),
                ])

            if rows:
                t = Table(rows, colWidths=[6.5 * cm, 11 * cm])
                t.setStyle(TableStyle([
                    ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("BACKGROUND", (0, 0), (0, -1), bg_section),
                    ("BOX", (0, 0), (-1, -1), 0.5, border),
                    ("INNERGRID", (0, 0), (-1, -1), 0.25, border),
                ]))
                block.append(t)
                block.append(Spacer(1, 0.2 * cm))

        story.append(KeepTogether(block[:3]) if len(block) > 3 else block[0])
        for item in block[1:]:
            story.append(item)
        story.append(PageBreak())

    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
    print(f"✓ PDF: {path}")
    return path


# ─── Fillable PDF (AcroForm) ─────────────────────────────────────────────────
def generate_fillable_pdf(output_path=None, preset_values=None, sector_label=None):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import cm

    preset_values = preset_values or {}
    path = output_path or (OUTPUT / "SwiftWithFer-Discovery-Document-Fillable.pdf")
    w, h = A4
    c = canvas.Canvas(str(path), pagesize=A4)
    c.setTitle(META["title"] + " (Compilabile)")

    primary = colors.HexColor(COLORS["primary"])
    dark = colors.HexColor(COLORS["dark"])
    text_light = colors.HexColor(COLORS["text_light"])
    border = colors.HexColor(COLORS["border"])

    page_num = 0
    y = h - 2.5 * cm
    margin_l = 2 * cm
    margin_r = w - 2 * cm
    field_w = margin_r - margin_l - 0.5 * cm
    line_h = 0.55 * cm

    def pdf_safe(text):
        """Sanitize text for AcroForm (Helvetica Standard encoding)."""
        replacements = {"€": "EUR ", "—": "-", "–": "-", """: "'", """: "'", "…": "..."}
        for old, new in replacements.items():
            text = text.replace(old, new)
        return text.encode("latin-1", errors="replace").decode("latin-1")

    def new_page():
        nonlocal y, page_num
        if page_num > 0:
            c.showPage()
        page_num += 1
        y = h - 2.5 * cm
        # Header
        c.setStrokeColor(primary)
        c.setLineWidth(1.5)
        c.line(margin_l, h - 1.5 * cm, margin_r, h - 1.5 * cm)
        if LOGO.exists():
            c.drawImage(str(LOGO), margin_l, h - 1.4 * cm, width=0.7 * cm, height=0.7 * cm, preserveAspectRatio=True, mask="auto")
        c.setFont("Helvetica", 7)
        c.setFillColor(text_light)
        c.drawString(margin_l + 1 * cm, h - 1.2 * cm, "SwiftWithFer — Discovery Document")
        c.drawRightString(margin_r, h - 1.2 * cm, META["confidentiality"])
        # Footer
        c.setStrokeColor(border)
        c.setLineWidth(0.5)
        c.line(margin_l, 1.5 * cm, margin_r, 1.5 * cm)
        c.setFont("Helvetica", 7)
        c.drawString(margin_l, 1 * cm, f"© {datetime.now().year} SwiftWithFer")
        c.drawRightString(margin_r, 1 * cm, f"Pagina {page_num}")

    def check_space(needed):
        nonlocal y
        if y - needed < 2.5 * cm:
            new_page()

    def draw_section_header(num, title, desc):
        nonlocal y
        check_space(2.5 * cm)
        c.setFillColor(primary)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(margin_l, y, f"{num}. {title}")
        y -= 0.5 * cm
        c.setFillColor(text_light)
        c.setFont("Helvetica-Oblique", 8)
        # Wrap description
        words = desc.split()
        line = ""
        for word in words:
            test = line + word + " "
            if c.stringWidth(test, "Helvetica-Oblique", 8) < field_w:
                line = test
            else:
                c.drawString(margin_l, y, line.strip())
                y -= 0.35 * cm
                line = word + " "
        if line:
            c.drawString(margin_l, y, line.strip())
            y -= 0.35 * cm
        c.setStrokeColor(primary)
        c.setLineWidth(0.5)
        c.line(margin_l, y, margin_r, y)
        y -= 0.4 * cm

    def draw_subsection(title):
        nonlocal y
        check_space(1 * cm)
        c.setFillColor(dark)
        c.setFont("Helvetica-Bold", 10)
        c.drawString(margin_l, y, title)
        y -= 0.5 * cm

    def add_text_field(field):
        nonlocal y
        label = field["label"] + (" *" if field.get("required") else "")
        ftype = field["type"]
        fid = field["id"]

        if ftype == "checkbox":
            check_space(0.6 * cm)
            c.acroForm.checkbox(name=fid, x=margin_l, y=y - 0.15 * cm, size=12, buttonStyle="check")
            c.setFont("Helvetica", 9)
            c.setFillColor(dark)
            c.drawString(margin_l + 0.5 * cm, y - 0.1 * cm, pdf_safe(label))
            y -= line_h
            return

        if ftype == "checkbox_group":
            check_space(0.5 * cm)
            c.setFont("Helvetica-Bold", 9)
            c.setFillColor(dark)
            c.drawString(margin_l, y, pdf_safe(label))
            y -= 0.4 * cm
            for i, opt in enumerate(field.get("options", [])):
                check_space(0.5 * cm)
                cb_name = f"{fid}_{i}"
                c.acroForm.checkbox(name=cb_name, x=margin_l + 0.3 * cm, y=y - 0.15 * cm, size=10, buttonStyle="check")
                c.setFont("Helvetica", 8)
                c.drawString(margin_l + 0.8 * cm, y - 0.1 * cm, pdf_safe(opt))
                y -= 0.4 * cm
            y -= 0.1 * cm
            return

        if ftype == "select":
            check_space(1.2 * cm)
            c.setFont("Helvetica", 8)
            c.setFillColor(dark)
            c.drawString(margin_l, y, pdf_safe(label))
            y -= 0.35 * cm
            opts = [pdf_safe(o) for o in field.get("options", [])]
            default = preset_values.get(fid)
            if default:
                default = pdf_safe(str(default))
            elif opts:
                default = opts[0]
            else:
                default = ""
            c.acroForm.choice(
                name=fid, x=margin_l, y=y - 0.5 * cm,
                width=field_w, height=0.55 * cm,
                value=default,
                options=opts,
                borderStyle="underlined",
                forceBorder=True,
            )
            y -= 0.9 * cm
            return

        # text, textarea, email, url, tel, number, date, currency
        height = 2.5 * cm if ftype == "textarea" else 0.55 * cm
        check_space(height + 0.6 * cm)
        c.setFont("Helvetica", 8)
        c.setFillColor(dark)
        c.drawString(margin_l, y, pdf_safe(label))
        y -= 0.35 * cm

        if ftype == "textarea":
            c.acroForm.textfield(
                name=fid, x=margin_l, y=y - height,
                width=field_w, height=height,
                borderStyle="underlined", forceBorder=True,
                fieldFlags="multiline",
                value=pdf_safe(str(preset_values.get(fid, ""))),
            )
            y -= height + 0.25 * cm
        else:
            c.acroForm.textfield(
                name=fid, x=margin_l, y=y - 0.5 * cm,
                width=field_w, height=0.55 * cm,
                borderStyle="underlined", forceBorder=True,
                value=pdf_safe(str(preset_values.get(fid, ""))),
            )
            y -= 0.9 * cm

    # Cover page
    new_page()
    if LOGO.exists():
        c.drawImage(str(LOGO), w / 2 - 1.5 * cm, h / 2 + 1 * cm, width=3 * cm, height=3 * cm, preserveAspectRatio=True, mask="auto")
    c.setFillColor(dark)
    c.setFont("Helvetica-Bold", 22)
    c.drawCentredString(w / 2, h / 2 - 0.5 * cm, "SwiftWithFer")
    c.setFont("Helvetica", 14)
    c.setFillColor(primary)
    c.drawCentredString(w / 2, h / 2 - 1.2 * cm, "Discovery Document")
    c.setFont("Helvetica", 10)
    c.setFillColor(text_light)
    c.drawCentredString(w / 2, h / 2 - 2 * cm, META["subtitle"])
    c.setFont("Helvetica", 9)
    cover_sub = f"Versione {META['version']} — Compilabile"
    if sector_label:
        cover_sub += f" — {pdf_safe(sector_label)}"
    c.drawCentredString(w / 2, h / 2 - 3.5 * cm, cover_sub)

    # Form pages
    for s in SECTIONS:
        new_page()
        draw_section_header(s["number"], s["title"], s["description"])
        for sub in s["subsections"]:
            draw_subsection(sub["title"])
            for field in sub["fields"]:
                add_text_field(field)

    c.save()
    print(f"✓ Fillable PDF: {path}")
    return path


def build_sector_json(sector_label: str) -> dict:
    """Return schema JSON with sector pre-filled."""
    schema = json.loads((OUTPUT / "SwiftWithFer-Discovery-Document.json").read_text(encoding="utf-8"))
    for section in schema["sections"]:
        for subsection in section["subsections"]:
            for field in subsection["fields"]:
                if field["id"] == "sector":
                    field["value"] = sector_label
    schema["meta"]["sector"] = sector_label
    return schema


def generate_sector_documents():
    """Generate fillable PDF per sector and deploy to public/."""
    import shutil

    docx_src = OUTPUT / "SwiftWithFer-Discovery-Document.docx"
    pdf_src = OUTPUT / "SwiftWithFer-Discovery-Document.pdf"

    PUBLIC_DISCOVERY.mkdir(parents=True, exist_ok=True)

    valid_ids = {sector["id"] for sector in SECTORS}
    for entry in PUBLIC_DISCOVERY.iterdir():
        if entry.is_dir() and entry.name not in valid_ids:
            shutil.rmtree(entry)

    for sector in SECTORS:
        sector_dir = PUBLIC_DISCOVERY / sector["id"]
        sector_dir.mkdir(parents=True, exist_ok=True)

        fillable_path = sector_dir / "SwiftWithFer-Discovery-Fillable.pdf"
        generate_fillable_pdf(
            output_path=fillable_path,
            preset_values={"sector": sector["label"]},
            sector_label=sector["label"],
        )

        if docx_src.exists():
            shutil.copy2(docx_src, sector_dir / "SwiftWithFer-Discovery.docx")
        if pdf_src.exists():
            shutil.copy2(pdf_src, sector_dir / "SwiftWithFer-Discovery.pdf")

        sector_json = build_sector_json(sector["label"])
        (sector_dir / "SwiftWithFer-Discovery.json").write_text(
            json.dumps(sector_json, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        print(f"✓ Sector: {sector['label']} → {sector_dir}")

    return PUBLIC_DISCOVERY


# ─── Main ────────────────────────────────────────────────────────────────────
def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    print(f"\n{'='*60}")
    print(f"  {META['title']}")
    print(f"  Generazione documenti in corso...")
    print(f"{'='*60}\n")

    generate_markdown()
    generate_json()
    generate_docx()
    generate_pdf()
    generate_fillable_pdf()
    generate_sector_documents()

    print(f"\n{'='*60}")
    print(f"  Completato! File in: {OUTPUT}")
    print(f"  Settori pubblicati in: {PUBLIC_DISCOVERY}")
    print(f"{'='*60}\n")

    # Stats
    total_fields = sum(1 for _ in iter_fields())
    print(f"  Sezioni: {len(SECTIONS)}")
    print(f"  Campi totali: {total_fields}")
    print()


if __name__ == "__main__":
    main()
